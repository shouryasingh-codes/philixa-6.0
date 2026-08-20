from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = Path('output/resume')
OUT.mkdir(parents=True, exist_ok=True)
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.42)
sec.bottom_margin = Inches(0.42)
sec.left_margin = Inches(0.55)
sec.right_margin = Inches(0.55)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Aptos'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos')
normal.font.size = Pt(8.8)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

accent = RGBColor(27, 75, 120)
dark = RGBColor(25, 25, 25)

def set_font(run, size=None, bold=None, color=None):
    run.font.name = 'Aptos'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos')
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = color

def para(text='', align=None, before=0, after=0, line=1.0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        set_font(r, 8.8, color=dark)
    return p

def section(title):
    p = para(before=4, after=1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title.upper())
    set_font(r, 9.4, True, accent)
    return p

def project(title, stack='', links=''):
    p = para(after=0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, 9.1, True, dark)
    if stack:
        r = p.add_run(' | ' + stack)
        set_font(r, 7.9, color=RGBColor(80,80,80))
    if links:
        r = p.add_run(' | ' + links)
        set_font(r, 7.9, color=accent)
    return p

def bullet(text):
    p = para(before=0, after=0, line=1.0)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.first_line_indent = Inches(-0.10)
    p.paragraph_format.keep_together = True
    r = p.add_run('• ' + text)
    set_font(r, 8.35, color=dark)
    return p

# Header
p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
r = p.add_run('SHOURYA SINGH')
set_font(r, 17, True, dark)
p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
r = p.add_run('Applied AI & Backend Developer')
set_font(r, 9.6, True, accent)
p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
r = p.add_run('+91 9610621152  |  shouryasingh3937@gmail.com  |  Jaipur, India  |  Open to Gurugram / Remote  |  LinkedIn  |  GitHub')
set_font(r, 8.2, color=dark)

section('Professional Summary')
p = para('Applied AI and backend developer building end-to-end LLM workflows, evidence-backed retrieval, voice-enabled copilots, and asynchronous Python services. Focused on reliable AI product features with validation, human review, and deployment-ready architecture.', after=1)
p.paragraph_format.keep_together = True

section('Projects')
project('PHILIXA - AI Meeting Intelligence MVP for Financial Advisors', '2026', 'Live Demo | GitHub')
bullet('Built an end-to-end meeting-intelligence workflow for pasted notes, audio uploads, and live browser microphone input; extracts commitments, risks, client priorities, and follow-up context.')
bullet('Implemented economy-model-first extraction with Pydantic validation, escalation to a review model, manual-review fallback, and audit logs for model latency and estimated cost.')
bullet('Built evidence-backed client memory using pgvector and Sentence Transformers; semantic retrieval returns meeting dates and snippets for natural-language client questions.')
bullet('Engineered browser voice workflows with AudioWorklet/WebSocket streaming, faster-whisper/Deepgram transcription, Web Speech API fast dictation, and a Deepgram STT + Sarvam TTS assistant for hands-free client queries.')
bullet('Added organization/user-scoped client data, HITL client resolution, and opt-in WhatsApp reminders with delivery logs, idempotent Redis/ARQ jobs, retries, quiet hours, and rate limits.')

project('Customer Retention AI System', '2025', 'Live Demo | GitHub')
bullet('Built and deployed a churn-prediction REST API using XGBoost, Scikit-learn, and SMOTE; returns churn probability, risk tier, and a suggested retention action through Swagger UI.')

project('SAVIOUR - AI Communication Simulator', '2026', 'GitHub')
bullet('Built a four-step communication-training workflow covering assessment, level detection, live role simulation, and a coaching report with AI-rewritten answers.')

section('Technical Skills')
p = para(after=0)
items = [
    ('Languages: ', 'Python, SQL, C++'),
    ('AI / LLM: ', 'LLM APIs, RAG, semantic search, provider routing, prompt engineering, Hinglish NLP'),
    ('Backend: ', 'FastAPI, Pydantic v2, SQLAlchemy ORM, REST APIs, WebSocket, Alembic'),
    ('Data & Async: ', 'PostgreSQL, pgvector, Redis, MinIO, ARQ job queues'),
    ('Voice / ML: ', 'faster-whisper, Deepgram, Sarvam TTS, pyannote.audio, AudioWorklet, Web Speech API, XGBoost, Scikit-learn, SMOTE, Sentence Transformers'),
    ('Delivery: ', 'Docker, Render, Git/GitHub, pytest, httpx'),
]
for label, value in items:
    r = p.add_run(label); set_font(r, 8.2, True, dark)
    r = p.add_run(value); set_font(r, 8.2, color=dark)
    p.add_run('\n')
p.paragraph_format.keep_together = True

section('Education')
p = para()
r = p.add_run('Bachelor of Computer Applications (BCA)')
set_font(r, 8.6, True, dark)
r = p.add_run(' | Suresh Gyan Vihar University, Jaipur | Expected 2027')
set_font(r, 8.6, color=dark)

doc.save(OUT / 'Shourya_Singh_Resume_Updated_2026.docx')
